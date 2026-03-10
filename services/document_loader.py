# from pypdf import PdfReader
# import docx

# def load_document(file):

#     if file.filename.endswith(".pdf"):
#         reader = PdfReader(file.file)
#         text = ""

#         for page in reader.pages:
#             text += page.extract_text()

#         return text

#     elif file.filename.endswith(".docx"):
#         doc = docx.Document(file.file)
#         text = "\n".join([para.text for para in doc.paragraphs])
#         return text

#     else:
#         return file.file.read().decode()
















import pdfplumber
import docx


def load_document(file):

    file_ext = file.name.split(".")[-1].lower()

    if file_ext == "pdf":

        text = ""

        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"

        return text

    elif file_ext == "docx":

        doc = docx.Document(file)

        text = "\n".join([p.text for p in doc.paragraphs])

        return text

    elif file_ext == "txt":

        return file.read().decode("utf-8")

    else:
        return ""
